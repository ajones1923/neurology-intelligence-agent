#!/usr/bin/env python3
"""Generate a DOCX report for the Neurology Intelligence Agent.

Produces a Word document summarizing the agent's capabilities, seed
data statistics, and collection configuration.

Usage:
    python scripts/generate_docx.py [--output report.docx]
"""

import argparse
import logging
import sys
from pathlib import Path

# Ensure project root on sys.path
PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
logger = logging.getLogger(__name__)


def generate_report(output_path: str = "neurology_agent_report.docx"):
    """Generate a DOCX report with agent summary information."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError:
        logger.error(
            "python-docx not installed. Install with: pip install python-docx"
        )
        # Fallback: generate a text file
        logger.info("Generating text report instead ...")
        _generate_text_report(output_path.replace(".docx", ".txt"))
        return

    from src.ingest.pubmed_neuro_parser import get_landmark_paper_count
    from src.ingest.neuroimaging_parser import get_neuroimaging_protocol_count
    from src.ingest.eeg_parser import get_eeg_pattern_count
    from src.models import NeuroWorkflowType, ClinicalScaleType

    doc = Document()
    doc.add_heading("Neurology Intelligence Agent Report", level=0)
    doc.add_paragraph(
        "This report summarizes the capabilities and seed data for the "
        "HCLS AI Factory Neurology Intelligence Agent."
    )

    # Seed Data Statistics
    doc.add_heading("Seed Data Statistics", level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Data Source"
    table.cell(0, 1).text = "Count"
    table.cell(1, 0).text = "Landmark Neurology Papers"
    table.cell(1, 1).text = str(get_landmark_paper_count())
    table.cell(2, 0).text = "Neuroimaging Protocols"
    table.cell(2, 1).text = str(get_neuroimaging_protocol_count())
    table.cell(3, 0).text = "EEG Patterns"
    table.cell(3, 1).text = str(get_eeg_pattern_count())

    # Workflows
    doc.add_heading("Clinical Workflows", level=1)
    for wf in NeuroWorkflowType:
        doc.add_paragraph(wf.value.replace("_", " ").title(), style="List Bullet")

    # Clinical Scales
    doc.add_heading("Clinical Scales", level=1)
    for scale in ClinicalScaleType:
        doc.add_paragraph(scale.value.upper(), style="List Bullet")

    doc.save(output_path)
    logger.info("DOCX report generated: %s", output_path)


def _generate_text_report(output_path: str):
    """Fallback text report when python-docx is not available."""
    from src.ingest.pubmed_neuro_parser import get_landmark_paper_count
    from src.ingest.neuroimaging_parser import get_neuroimaging_protocol_count
    from src.ingest.eeg_parser import get_eeg_pattern_count
    from src.models import NeuroWorkflowType, ClinicalScaleType

    lines = [
        "Neurology Intelligence Agent Report",
        "=" * 40,
        "",
        "Seed Data Statistics:",
        f"  Landmark Papers: {get_landmark_paper_count()}",
        f"  Neuroimaging Protocols: {get_neuroimaging_protocol_count()}",
        f"  EEG Patterns: {get_eeg_pattern_count()}",
        "",
        "Clinical Workflows:",
    ]
    for wf in NeuroWorkflowType:
        lines.append(f"  - {wf.value.replace('_', ' ').title()}")

    lines.append("")
    lines.append("Clinical Scales:")
    for scale in ClinicalScaleType:
        lines.append(f"  - {scale.value.upper()}")

    Path(output_path).write_text("\n".join(lines))
    logger.info("Text report generated: %s", output_path)


def main():
    """Parse arguments and generate report."""
    parser = argparse.ArgumentParser(description="Generate neurology agent DOCX report")
    parser.add_argument(
        "--output",
        default="neurology_agent_report.docx",
        help="Output file path (default: neurology_agent_report.docx)",
    )
    args = parser.parse_args()
    generate_report(args.output)


if __name__ == "__main__":
    main()
