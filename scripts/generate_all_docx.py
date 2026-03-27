#!/usr/bin/env python3
"""Generate DOCX versions of all documentation markdown files."""
import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = PROJECT_ROOT / "docs"

try:
    from docx import Document
except ImportError:
    print("python-docx not installed. pip install python-docx")
    sys.exit(1)

MD_FILES = [
    "PRODUCTION_READINESS_REPORT",
    "PROJECT_BIBLE",
    "ARCHITECTURE_GUIDE",
    "WHITE_PAPER",
    "DEPLOYMENT_GUIDE",
    "DEMO_GUIDE",
    "LEARNING_GUIDE_FOUNDATIONS",
    "LEARNING_GUIDE_ADVANCED",
    "INDEX",
]

for name in MD_FILES:
    md_path = DOCS_DIR / f"{name}.md"
    docx_path = DOCS_DIR / f"{name}.docx"
    if not md_path.exists():
        print(f"SKIP: {md_path} not found")
        continue
    text = md_path.read_text()
    doc = Document()
    for line in text.split("\n"):
        s = line.strip()
        if s.startswith("# "):
            doc.add_heading(s[2:], level=0)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=1)
        elif s.startswith("### "):
            doc.add_heading(s[4:], level=2)
        elif s.startswith("#### "):
            doc.add_heading(s[5:], level=3)
        elif s.startswith("- ") or s.startswith("* "):
            doc.add_paragraph(s[2:], style="List Bullet")
        elif s == "" or s.startswith("---"):
            continue
        else:
            doc.add_paragraph(s)
    doc.save(str(docx_path))
    print(f"OK: {docx_path}")

print("Done generating DOCX files.")
